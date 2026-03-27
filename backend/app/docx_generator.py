"""Word (.docx) resume generator — zero external binary dependencies."""
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.app.config import RESUMES_DIR

_ACCENT = RGBColor(30, 90, 160)
_FONT = "Calibri"


def _set_font(run, size: int, bold: bool = False, italic: bool = False, color: RGBColor | None = None):
    run.font.name = _FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_bottom_border(paragraph):
    """Add a thin bottom border line to a paragraph (section heading rule)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E5AA0")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc: Document, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    _set_font(run, 12, bold=True, color=_ACCENT)
    _add_bottom_border(p)
    return p


def _bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    _set_font(run, 10)
    return p


def _build_document(
    name: str,
    summary: str,
    skills: list[str],
    experience: list[dict],
    projects: list[dict],
    education: list,
) -> Document:
    doc = Document()

    # Margins — 1 inch all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Name header
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(name)
    _set_font(name_run, 20, bold=True, color=_ACCENT)

    # Summary
    if summary:
        _section_heading(doc, "Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(summary)
        _set_font(run, 10)

    # Skills
    if skills:
        _section_heading(doc, "Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("  |  ".join(skills))
        _set_font(run, 10)

    # Experience
    if experience:
        _section_heading(doc, "Experience")
        for exp in experience:
            role = exp.get("role", "")
            company = exp.get("company", "")
            duration = exp.get("duration", "")
            bullets = exp.get("bullets", [])

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            title_run = p.add_run(f"{role} @ {company}")
            _set_font(title_run, 11, bold=True)

            if duration:
                dur_p = doc.add_paragraph()
                dur_p.paragraph_format.space_before = Pt(0)
                dur_p.paragraph_format.space_after = Pt(1)
                dur_run = dur_p.add_run(duration)
                _set_font(dur_run, 9, italic=True)

            for b in bullets:
                text = b if isinstance(b, str) else b.get("raw", b.get("rewritten", ""))
                if text:
                    _bullet(doc, text)

    # Projects
    if projects:
        _section_heading(doc, "Projects")
        for proj in projects:
            proj_name = proj.get("name", "")
            bullets = proj.get("bullets", [])

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            name_run = p.add_run(proj_name)
            _set_font(name_run, 11, bold=True)

            for b in bullets:
                if isinstance(b, dict):
                    text = b.get("rewritten") or b.get("raw", "")
                else:
                    text = str(b)
                if text:
                    _bullet(doc, text)

    # Education
    if education:
        _section_heading(doc, "Education")
        for edu in education:
            if hasattr(edu, "degree"):
                degree = edu.degree
                field = getattr(edu, "field", "")
                institution = getattr(edu, "institution", "")
                duration = getattr(edu, "duration", "")
            else:
                degree = edu.get("degree", "")
                field = edu.get("field", "")
                institution = edu.get("institution", "")
                duration = edu.get("duration", "")

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            degree_text = f"{degree} in {field}" if field else degree
            deg_run = p.add_run(degree_text)
            _set_font(deg_run, 11, bold=True)

            inst_text = f"{institution} · {duration}" if duration else institution
            if inst_text:
                inst_p = doc.add_paragraph()
                inst_p.paragraph_format.space_before = Pt(0)
                inst_p.paragraph_format.space_after = Pt(1)
                inst_run = inst_p.add_run(inst_text)
                _set_font(inst_run, 10, italic=True)

    return doc


def generate_base_resume(profile) -> Path:
    """Build base Word resume from full profile. Saved to data/resumes/base_resume.docx."""
    RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    output_path = RESUMES_DIR / "base_resume.docx"

    experience = [
        {
            "role": exp.role,
            "company": exp.company,
            "duration": exp.duration,
            "bullets": [b.raw for b in exp.bullets],
        }
        for exp in profile.experience
    ]

    projects = [
        {
            "name": proj.name,
            "bullets": [{"raw": b.raw} for b in proj.bullets],
        }
        for proj in profile.projects
    ]

    doc = _build_document(
        name=profile.name,
        summary="",
        skills=[s.name for s in profile.skills],
        experience=experience,
        projects=projects,
        education=profile.education,
    )
    doc.save(output_path)
    return output_path


def generate_tailored_resume(tailored_content: dict, profile, output_path: Path) -> str:
    """Build tailored Word resume. Combines tailored_content with profile.education."""
    RESUMES_DIR.mkdir(parents=True, exist_ok=True)

    doc = _build_document(
        name=tailored_content.get("name") or profile.name,
        summary=tailored_content.get("summary", ""),
        skills=tailored_content.get("skills", []),
        experience=tailored_content.get("experience", []),
        projects=tailored_content.get("projects", []),
        education=getattr(profile, "education", []),
    )
    doc.save(output_path)
    return str(output_path)
